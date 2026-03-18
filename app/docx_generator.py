"""
Google Docs–compatible report generator — converts Markdown SEO reports to .docx.

Pipeline: Markdown → parsed sections → styled python-docx document.

Output .docx files open natively in Google Docs when uploaded to Drive.
"""

import re
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ---------------------------------------------------------------------------
# Color palette (matches our PDF theme)
# ---------------------------------------------------------------------------
BLUE_DARK = RGBColor(0x1E, 0x3A, 0x5F)
BLUE_ACCENT = RGBColor(0x25, 0x63, 0xEB)
GRAY_TEXT = RGBColor(0x55, 0x55, 0x55)
GRAY_LIGHT = RGBColor(0x64, 0x74, 0x8B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
ZEBRA_BG = "F8FAFC"
HEADER_BG = "1E3A5F"
BORDER_COLOR = "CBD5E1"
PROFILE_BG = "F0F9FF"
PROFILE_BORDER = "BAE6FD"


def _set_cell_bg(cell, color_hex: str):
    """Set background color on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_borders(cell, color: str = BORDER_COLOR, sz: int = 4):
    """Set thin borders on a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def _add_styled_table(doc, headers: list[str], rows: list[list[str]]):
    """Add a professionally styled table to the document."""
    if not headers:
        return

    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    hdr_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header_text.strip())
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = WHITE
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_cell_bg(cell, HEADER_BG)
        _set_cell_borders(cell, HEADER_BG)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[1 + row_idx]
        for col_idx in range(num_cols):
            cell = row.cells[col_idx]
            text = row_data[col_idx].strip() if col_idx < len(row_data) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            # Handle bold markers in cell text
            _add_rich_text(p, text, base_size=Pt(8))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_cell_borders(cell)
            # Zebra striping
            if row_idx % 2 == 1:
                _set_cell_bg(cell, ZEBRA_BG)

    # Compact spacing
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(1)
                paragraph.paragraph_format.space_after = Pt(1)

    doc.add_paragraph("")  # spacing after table


def _add_rich_text(paragraph, text: str, base_size=Pt(10), base_color=None):
    """Add text with **bold**, *italic*, and <br> line break support."""
    from docx.oxml.ns import qn as _qn
    # First split on <br> tags to handle line breaks
    br_parts = re.split(r'<br\s*/?>', text)
    for br_idx, br_part in enumerate(br_parts):
        if br_idx > 0:
            # Add a line break
            run = paragraph.add_run()
            run.font.size = base_size
            run.font.name = "Calibri"
            run._r.append(parse_xml(f'<w:br {nsdecls("w")}/>'))
        # Split on bold and italic markers
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', br_part)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                run = paragraph.add_run(part)
            run.font.size = base_size
            run.font.name = "Calibri"
            if base_color:
                run.font.color.rgb = base_color


def _parse_md_table(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    """Parse a markdown table into headers and rows."""
    if len(lines) < 2:
        return [], []

    def split_row(line):
        line = line.strip()
        if line.startswith("|"):
            line = line[1:]
        if line.endswith("|"):
            line = line[:-1]
        return [c.strip() for c in line.split("|")]

    headers = split_row(lines[0])
    # Skip separator line (line[1])
    rows = []
    for line in lines[2:]:
        if line.strip() and not re.match(r'^\|?\s*[-:]+', line):
            rows.append(split_row(line))
    return headers, rows


def _extract_header_info(md_content: str) -> dict:
    """Extract report title, URL, keyword, and date from the markdown header."""
    info = {"title": "", "url": "", "keyword": "", "date": ""}

    h1 = re.search(r'^#\s+(.+)$', md_content, re.MULTILINE)
    if h1:
        info["title"] = h1.group(1).strip()

    # New format: **URL:** and **Keyword:** on separate lines
    url_match = re.search(r'^\*\*URL:\*\*\s*(.+)$', md_content, re.MULTILINE)
    if url_match:
        info["url"] = url_match.group(1).strip()

    kw_match = re.search(r'^\*\*Keyword:\*\*\s*(.+)$', md_content, re.MULTILINE)
    if kw_match:
        info["keyword"] = kw_match.group(1).strip()

    # Legacy format: ## Target: URL | Keyword: KW
    if not info["url"]:
        target = re.search(r'^##\s+Target:\s*(.*?)\s*\|\s*Keyword:\s*(.+)$', md_content, re.MULTILINE)
        if target:
            info["url"] = target.group(1).strip()
            info["keyword"] = target.group(2).strip()

    date = re.search(r'^\*Generated:\s*(.+?)\*$', md_content, re.MULTILINE)
    if date:
        info["date"] = date.group(1).strip()

    return info


def generate_docx(md_content: str, output_path: str) -> dict:
    """
    Convert a Markdown SEO report to a styled .docx (Google Docs compatible).

    Args:
        md_content: The full markdown report content.
        output_path: Path to save the .docx file.

    Returns:
        {"success": bool, "error": str|None, "path": str}
    """
    try:
        doc = Document()

        # Page setup: A4, reasonable margins
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

        # Default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(10)
        font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

        # Extract header info
        header_info = _extract_header_info(md_content)

        # --- Report Header (styled block) ---
        _add_report_header(doc, header_info)

        # --- Parse and render body ---
        # Remove the header lines (title, target, date, first hr)
        body = md_content
        body = re.sub(r'^#\s+.+\n', '', body, count=1)
        body = re.sub(r'^##\s+Target:.+\n', '', body, count=1)
        body = re.sub(r'^\*Generated:.+\*\n?', '', body, count=1)
        body = re.sub(r'^---\s*\n?', '', body.lstrip(), count=1)

        _render_body(doc, body.strip())

        # Save
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))

        return {"success": True, "error": None, "path": str(out)}

    except Exception as e:
        return {"success": False, "error": str(e), "path": ""}


def _add_report_header(doc, info: dict):
    """Add a styled report header block."""
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_p.paragraph_format.space_after = Pt(4)
    run = title_p.add_run(info.get("title") or "SEO Optimization Report")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = BLUE_DARK
    run.font.name = "Calibri"

    # Subtitle line with URL + keyword
    if info.get("url") or info.get("keyword"):
        sub_p = doc.add_paragraph()
        sub_p.paragraph_format.space_before = Pt(0)
        sub_p.paragraph_format.space_after = Pt(2)
        if info.get("url"):
            run = sub_p.add_run("URL: ")
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = GRAY_LIGHT
            run.font.name = "Calibri"
            run = sub_p.add_run(info["url"])
            run.font.size = Pt(10)
            run.font.color.rgb = BLUE_ACCENT
            run.font.name = "Calibri"

        kw_p = doc.add_paragraph()
        kw_p.paragraph_format.space_before = Pt(0)
        kw_p.paragraph_format.space_after = Pt(2)
        if info.get("keyword"):
            run = kw_p.add_run("Keyword: ")
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = GRAY_LIGHT
            run.font.name = "Calibri"
            run = kw_p.add_run(info["keyword"])
            run.font.size = Pt(10)
            run.font.color.rgb = BLUE_DARK
            run.font.name = "Calibri"

    if info.get("date"):
        date_p = doc.add_paragraph()
        date_p.paragraph_format.space_before = Pt(0)
        date_p.paragraph_format.space_after = Pt(6)
        run = date_p.add_run(f"Generated: {info['date']}")
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY_TEXT
        run.font.name = "Calibri"
        run.italic = True

    # Divider line
    border_p = doc.add_paragraph()
    border_p.paragraph_format.space_before = Pt(4)
    border_p.paragraph_format.space_after = Pt(8)
    pPr = border_p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:color="2563EB"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def _render_body(doc, body: str):
    """Parse markdown body and render into the document."""
    lines = body.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^---+\s*$', stripped):
            _add_hr(doc)
            i += 1
            continue

        # Headings
        h_match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
        if h_match:
            level = len(h_match.group(1))
            text = h_match.group(2).strip()
            _add_heading(doc, text, level)
            i += 1
            continue

        # Table detection — separator must look like |---|---| (not just start with -)
        if "|" in stripped and i + 1 < len(lines) and re.match(r'^\|[\s\-:|]+\|$', lines[i + 1].strip()):
            table_lines = []
            while i < len(lines) and "|" in lines[i].strip():
                table_lines.append(lines[i])
                i += 1
            headers, rows = _parse_md_table(table_lines)
            if headers and rows:
                _add_styled_table(doc, headers, rows)
            continue

        # Bullet list
        if re.match(r'^[\-\*]\s+', stripped):
            list_items = []
            while i < len(lines) and re.match(r'^\s*[\-\*]\s+', lines[i]):
                item_text = re.sub(r'^\s*[\-\*]\s+', '', lines[i]).strip()
                list_items.append(item_text)
                i += 1
            for item in list_items:
                _add_bullet(doc, item)
            continue

        # Numbered list
        if re.match(r'^\d+[\.\)]\s+', stripped):
            list_items = []
            while i < len(lines) and re.match(r'^\s*\d+[\.\)]\s+', lines[i]):
                item_text = re.sub(r'^\s*\d+[\.\)]\s+', '', lines[i]).strip()
                list_items.append(item_text)
                i += 1
            for idx, item in enumerate(list_items, 1):
                _add_numbered_item(doc, item, idx)
            continue

        # Regular paragraph
        para_lines = []
        while i < len(lines) and lines[i].strip() and not re.match(r'^(#{1,4}\s|---|\||\-\s|\*\s|\d+[\.\)])', lines[i].strip()):
            para_lines.append(lines[i].strip())
            i += 1
        if para_lines:
            text = " ".join(para_lines)
            _add_paragraph(doc, text)
        else:
            i += 1


def _add_heading(doc, text: str, level: int):
    """Add a styled heading using Word's built-in Heading styles for navigation pane support."""
    # Clean markdown formatting from heading text
    clean = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    clean = re.sub(r'\*(.*?)\*', r'\1', clean)

    # Use built-in Heading styles so the navigation pane picks them up
    heading_level = min(level, 4)  # Word supports Heading 1-9
    p = doc.add_heading(clean, level=heading_level)
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True

    # Override the built-in heading style with our custom look
    for run in p.runs:
        run.font.name = "Calibri"
        if level == 1:
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = BLUE_DARK
        elif level == 2:
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = BLUE_DARK
        elif level == 3:
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        elif level == 4:
            run.bold = True
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    # Add bottom border for H2
    if level == 2:
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="4" w:color="{BORDER_COLOR}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)


def _add_paragraph(doc, text: str):
    """Add a paragraph with inline markdown formatting."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    _add_rich_text(p, text, base_size=Pt(10))


def _add_bullet(doc, text: str):
    """Add a bullet list item."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    # Clear auto-generated text, use rich formatting
    p.clear()
    _add_rich_text(p, text, base_size=Pt(10))


def _add_numbered_item(doc, text: str, number: int):
    """Add a numbered list item."""
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.clear()
    _add_rich_text(p, text, base_size=Pt(10))


def _add_hr(doc):
    """Add a horizontal rule (bottom border on empty paragraph)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:color="E2E8F0"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def md_file_to_docx(md_path: str, docx_path: Optional[str] = None) -> dict:
    """
    Convenience: read a .md file and convert to .docx.

    Args:
        md_path: Path to the markdown file.
        docx_path: Output .docx path. If None, uses same name with .docx extension.

    Returns:
        {"success": bool, "error": str|None, "path": str}
    """
    md_file = Path(md_path)
    if not md_file.exists():
        return {"success": False, "error": f"File not found: {md_path}", "path": ""}

    md_content = md_file.read_text(encoding="utf-8")

    if not docx_path:
        docx_path = str(md_file.with_suffix(".docx"))

    return generate_docx(md_content, docx_path)
